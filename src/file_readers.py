# src/file_readers.py
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
import os
from io import BytesIO, StringIO

def read_txt(file) -> str:
    """Read TXT from path or file-like object."""
    if hasattr(file, "read"):  # file-like object
        file.seek(0)
        return file.read().decode("utf-8", errors="ignore")
    else:  # path string
        with open(file, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

def read_pdf(file) -> str:
    """Read PDF from path or file-like object."""
    if hasattr(file, "read"):
        file.seek(0)
        # pdfminer can read a BytesIO stream
        return pdf_extract_text(file) or ""
    else:
        return pdf_extract_text(file) or ""

def read_docx(file) -> str:
    """Read DOCX from path or file-like object."""
    if hasattr(file, "read"):
        file.seek(0)
        doc = Document(BytesIO(file.read()))
    else:
        doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs)

def read_any(file) -> str:
    """Detect file type and read accordingly."""
    # Determine extension
    if hasattr(file, "name"):  # UploadedFile has .name attribute
        ext = os.path.splitext(file.name)[1].lower()
    elif isinstance(file, str):
        ext = os.path.splitext(file)[1].lower()
    else:
        raise ValueError("Cannot determine file type.")

    if ext == ".pdf":
        return read_pdf(file)
    elif ext == ".docx":
        return read_docx(file)
    elif ext == ".txt":
        return read_txt(file)
    else:
        raise ValueError("Unsupported file type. Use PDF, DOCX, or TXT.")
